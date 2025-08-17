"""
Financial RAG System for TradingAgents-CN

This module implements a Retrieval-Augmented Generation (RAG) system specifically
designed for financial analysis with domain-specific knowledge base, vector embeddings,
and semantic search capabilities.
"""

import hashlib
import os
import pickle
from collections import defaultdict
from dataclasses import dataclass, field
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any

import numpy as np
import pandas as pd
import requests

# Vector database and embeddings
try:
    import chromadb
    from chromadb.config import Settings

    CHROMADB_AVAILABLE = True
except ImportError:
    CHROMADB_AVAILABLE = False
    print("ChromaDB not available - RAG system will use fallback storage")

try:
    from sentence_transformers import SentenceTransformer  # noqa: F401

    SENTENCE_TRANSFORMERS_AVAILABLE = True
except ImportError:
    SENTENCE_TRANSFORMERS_AVAILABLE = False
    print("Sentence Transformers not available - using fallback embedding method")

# TradingAgents imports
from tradingagents.dataflows import get_finnhub_news, get_YFin_data_window
from tradingagents.utils.logging_init import get_logger

logger = get_logger("financial_rag")


@dataclass
class FinancialDocument:
    """Financial document structure for RAG system"""

    doc_id: str
    title: str
    content: str
    doc_type: str  # "news", "research", "earnings", "filing", "market_data"
    symbol: str | None = None
    sector: str | None = None
    market: str | None = None
    timestamp: datetime = field(default_factory=datetime.now)
    metadata: dict[str, Any] = field(default_factory=dict)
    embedding: np.ndarray | None = None
    relevance_score: float = 0.0


@dataclass
class RAGQuery:
    """RAG query structure"""

    query_text: str
    query_type: str = "general"  # "general", "technical", "fundamental", "news", "risk"
    symbols: list[str] | None = None
    user_id: str | None = None
    date_range: tuple[datetime, datetime] | None = None
    doc_types: list[str] | None = None
    top_k: int = 5
    relevance_threshold: float = 0.7
    context: dict[str, Any] = field(default_factory=dict)


@dataclass
class RAGResponse:
    """RAG response structure"""

    query: RAGQuery
    retrieved_documents: list[FinancialDocument]
    generated_response: str
    confidence_score: float
    sources: list[str] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)


class FinancialEmbedding:
    """Financial domain-specific embedding system"""

    def __init__(self, model_name: str = "all-MiniLM-L6-v2"):
        self.model_name = model_name
        self.model = None
        # provider & remote config
        # Always use remote API embedding by default (Qwen3-Embedding-8B via SiliconFlow)
        self.provider = os.getenv("EMBEDDING_PROVIDER", "siliconflow").strip().lower()
        self.remote_model = os.getenv("EMBEDDING_MODEL", "Qwen/Qwen3-Embedding-8B")
        try:
            self.remote_dim = int(os.getenv("EMBEDDING_DIM", "4096"))
        except Exception:
            self.remote_dim = None
        self.siliconflow_api_key = os.getenv("SILICONFLOW_API_KEY")
        self.siliconflow_base = (
            os.getenv("SILICONFLOW_BASE_URL") or "https://api.siliconflow.cn/v1"
        ).rstrip("/")

        # Financial domain keywords and weights
        self.financial_keywords = {
            "technical": [
                "support",
                "resistance",
                "trend",
                "moving average",
                "rsi",
                "macd",
                "volume",
                "breakout",
                "momentum",
                "volatility",
            ],
            "fundamental": [
                "earnings",
                "revenue",
                "profit",
                "cash flow",
                "debt",
                "ratio",
                "growth",
                "valuation",
                "dividend",
                "balance sheet",
            ],
            "risk": [
                "risk",
                "volatility",
                "downside",
                "correlation",
                "var",
                "drawdown",
                "stress test",
                "scenario",
                "hedge",
                "exposure",
            ],
            "market": [
                "bull",
                "bear",
                "sentiment",
                "market",
                "sector",
                "industry",
                "economic",
                "federal reserve",
                "inflation",
                "gdp",
            ],
        }

        self._initialize_model()

    def _initialize_model(self):
        """Initialize embedding model/provider.

        Priority:
          1) If provider=siliconflow or (auto and no local model available but SiliconFlow configured) -> use SiliconFlow embeddings.
          2) If sentence_transformers available -> use local model.
          3) Fallback hash-based embedding.
        """
        try:
            # Force remote usage; local models are not used by design per project policy
            self.model = None  # indicate remote provider
            if self.siliconflow_api_key:
                logger.info(
                    f"Using SiliconFlow embeddings (forced): model={self.remote_model}, dim={self.remote_dim}"
                )
            else:
                logger.warning(
                    "SiliconFlow API key not set; embedding will fallback to hashing"
                )
        except Exception as e:
            logger.error(f"Failed to initialize embedding provider: {e}")
            self.model = None

    def embed_text(self, text: str, doc_type: str = "general") -> np.ndarray:
        """
        Create embeddings with financial domain enhancement

        Args:
            text: Text to embed
            doc_type: Type of document for domain weighting

        Returns:
            np.ndarray: Embedding vector
        """
        try:
            # Always prefer remote SiliconFlow embeddings via OpenAI-compatible API
            if self.siliconflow_api_key:
                emb = self._embed_via_openai_compat([text])
                if emb is not None:
                    base_embedding = np.array(emb[0], dtype=float)
                    enhanced_embedding = self._apply_domain_weighting(
                        text, base_embedding, doc_type
                    )
                    return enhanced_embedding
            # Fallback: hashing (when API not configured or call fails)
            return self._fallback_embedding(text)
        except Exception as e:
            logger.error(f"Embedding failed: {e}")
            return self._fallback_embedding(text)

    def _embed_via_openai_compat(self, inputs: list[str]) -> list[list[float]] | None:
        """Call OpenAI-compatible embeddings endpoint (SiliconFlow) and return vectors.

        Returns list of embeddings for each input text, or None on failure.
        """
        try:
            url = f"{self.siliconflow_base}/embeddings"
            headers = {
                "Authorization": f"Bearer {self.siliconflow_api_key}",
                "Content-Type": "application/json",
            }
            payload: dict[str, Any] = {
                "model": self.remote_model,
                "input": inputs,
            }
            if self.remote_dim and isinstance(self.remote_dim, int):
                payload["dimensions"] = self.remote_dim
            resp = requests.post(url, headers=headers, json=payload, timeout=30)
            if resp.status_code >= 400:
                logger.warning(
                    f"SiliconFlow embeddings error {resp.status_code}: {resp.text[:200]}"
                )
                return None
            data = resp.json()
            if not data or "data" not in data:
                return None
            # data.data is list of objects with 'embedding'
            vectors = [item.get("embedding") for item in data["data"]]
            if any(v is None for v in vectors):
                return None
            return vectors
        except Exception as e:
            logger.warning(f"OpenAI-compatible embedding call failed: {e}")
            return None

    def _apply_domain_weighting(
        self, text: str, base_embedding: np.ndarray, doc_type: str
    ) -> np.ndarray:
        """Apply financial domain-specific weighting to embeddings"""
        # Calculate domain relevance scores
        text_lower = text.lower()
        domain_scores = {}

        for domain, keywords in self.financial_keywords.items():
            score = sum(1 for keyword in keywords if keyword in text_lower)
            domain_scores[domain] = score / len(keywords)  # Normalize by keyword count

        # Apply weighting based on document type
        weight_multiplier = 1.0
        if doc_type == "technical" and domain_scores.get("technical", 0) > 0.1:
            weight_multiplier = 1.2
        elif doc_type == "fundamental" and domain_scores.get("fundamental", 0) > 0.1:
            weight_multiplier = 1.2
        elif doc_type == "news" and domain_scores.get("market", 0) > 0.1:
            weight_multiplier = 1.1

        # Apply subtle enhancement to preserve semantic meaning
        enhanced_embedding = base_embedding * weight_multiplier

        # Normalize to maintain vector properties
        return enhanced_embedding / np.linalg.norm(enhanced_embedding)

    def _fallback_embedding(self, text: str) -> np.ndarray:
        """Fallback embedding method using simple hashing.

        Pads/truncates to the target dimension to avoid Chroma dimension mismatch.
        """
        # Simple hash-based embedding as fallback
        text_hash = hashlib.md5(text.encode()).hexdigest()

        # Base vector from hash (up to 32 dims)
        base = np.array(
            [
                int(text_hash[i : i + 2], 16) / 255.0
                for i in range(0, min(len(text_hash), 64), 2)
            ]
        )

        # Target dimension: use remote_dim if set, else 32
        target_dim = int(self.remote_dim) if self.remote_dim else 32
        if target_dim <= 0:
            target_dim = 32

        if len(base) < target_dim:
            embedding = np.pad(base, (0, target_dim - len(base)))
        else:
            embedding = base[:target_dim]

        return embedding

    def compute_similarity(
        self, embedding1: np.ndarray, embedding2: np.ndarray
    ) -> float:
        """Compute cosine similarity between embeddings"""
        try:
            # Cosine similarity
            dot_product = np.dot(embedding1, embedding2)
            norm_product = np.linalg.norm(embedding1) * np.linalg.norm(embedding2)

            if norm_product == 0:
                return 0.0

            similarity = dot_product / norm_product
            return float(np.clip(similarity, -1.0, 1.0))

        except Exception as e:
            logger.error(f"Similarity computation failed: {e}")
            return 0.0


class FinancialKnowledgeBase:
    """Financial knowledge base with vector storage"""

    def __init__(
        self,
        storage_path: str = "financial_kb",
        pipeline: str | None = None,
        embedding_config: dict[str, Any] | None = None,
    ):
        self.storage_path = Path(storage_path)
        self.storage_path.mkdir(parents=True, exist_ok=True)

        # Pipeline selection: vector | graphrag | raptor (vector by default)
        try:
            import os as _os

            self.pipeline = (
                (pipeline or _os.getenv("RAG_PIPELINE") or "vector").strip().lower()
            )
            if self.pipeline not in {"vector", "graphrag", "raptor"}:
                self.pipeline = "vector"
        except Exception:
            self.pipeline = "vector"

        # Initialize vector database
        self.vector_db = None
        self.collection = None
        self._initialize_vector_db()

        # Initialize embedding system
        self.embedding_system = FinancialEmbedding()
        # Apply overrides if provided
        if embedding_config:
            try:
                if "provider" in embedding_config and embedding_config["provider"]:
                    self.embedding_system.provider = (
                        str(embedding_config["provider"]).strip().lower()
                    )
                if "model" in embedding_config and embedding_config["model"]:
                    self.embedding_system.remote_model = str(embedding_config["model"])
                if "dim" in embedding_config and embedding_config["dim"]:
                    self.embedding_system.remote_dim = int(embedding_config["dim"])
                if "api_key" in embedding_config and embedding_config["api_key"]:
                    self.embedding_system.siliconflow_api_key = str(
                        embedding_config["api_key"]
                    )
                if "base_url" in embedding_config and embedding_config["base_url"]:
                    self.embedding_system.siliconflow_base = str(
                        embedding_config["base_url"]
                    ).rstrip("/")
                logger.info(
                    f"Embedding overrides applied: provider={self.embedding_system.provider}, model={self.embedding_system.remote_model}, dim={self.embedding_system.remote_dim}"
                )
            except Exception as e:
                logger.warning(f"Failed to apply embedding overrides: {e}")

        # In-memory indexes for fast retrieval
        self.symbol_index: dict[str, list[str]] = defaultdict(list)
        self.date_index: dict[str, list[str]] = defaultdict(list)
        self.type_index: dict[str, list[str]] = defaultdict(list)

        # Document storage
        self.documents: dict[str, FinancialDocument] = {}

        # Load existing knowledge base
        self._load_knowledge_base()

        logger.info("Financial Knowledge Base initialized")

    def build_evidence_map(self, documents: list[FinancialDocument]) -> dict[str, Any]:
        """Lightweight evidence map for compatibility with GraphRAG/RAPTOR plans.

        This does not implement full graph or hierarchy; it extracts simple
        entity-like tokens and groups by doc_type/source for now.
        """
        evidence: dict[str, Any] = {"entities": {}, "sources": []}
        try:
            for doc in documents:
                # naive token extraction of uppercase words and symbols
                import re

                ents = re.findall(r"\b[A-Z][A-Z0-9_\-]{1,}\b", doc.content or "")
                for e in ents[:20]:  # cap per doc
                    evidence["entities"].setdefault(e, 0)
                    evidence["entities"][e] += 1
                if doc.metadata.get("url"):
                    evidence["sources"].append(doc.metadata["url"])
        except Exception:
            pass
        return evidence

    def _initialize_vector_db(self):
        """Initialize ChromaDB vector database"""
        try:
            # Allow disabling vector DB via env (fallback to in-memory similarity)
            mem_flag = os.getenv("MEMORY_ENABLED", "true").strip().lower()
            if mem_flag in {"0", "false", "no", "off"}:
                logger.warning("Vector DB disabled by MEMORY_ENABLED=false; using in-memory search")
                self.vector_db = None
                self.collection = None
                return

            if CHROMADB_AVAILABLE:
                # Initialize persistent ChromaDB client (Windows/WSL friendly)
                db_path = str(self.storage_path / "chromadb")
                client = None
                try:
                    from tradingagents.agents.utils.chromadb_win11_config import (
                        get_optimal_chromadb_client,
                    )

                    client = get_optimal_chromadb_client(db_path)
                except Exception:
                    client = chromadb.PersistentClient(
                        path=db_path,
                        settings=Settings(
                            anonymized_telemetry=False, allow_reset=True
                        ),
                    )

                self.vector_db = client
                # Create or get collection
                self.collection = self.vector_db.get_or_create_collection(
                    name="financial_knowledge",
                    metadata={"description": "Financial documents and data"},
                )

                logger.info("ChromaDB initialized successfully")
            else:
                logger.warning("ChromaDB not available - using fallback storage")

        except Exception as e:
            logger.error(f"Failed to initialize vector database: {e}")
            self.vector_db = None
            self.collection = None

    def add_document(self, document: FinancialDocument) -> bool:
        """
        Add document to knowledge base

        Args:
            document: Financial document to add

        Returns:
            bool: Success status
        """
        try:
            # Generate embedding if not provided
            if document.embedding is None:
                document.embedding = self.embedding_system.embed_text(
                    document.content, document.doc_type
                )

            # Store in vector database
            if self.collection is not None:
                self.collection.add(
                    embeddings=[document.embedding.tolist()],
                    documents=[document.content],
                    metadatas=[
                        {
                            "doc_id": document.doc_id,
                            "title": document.title,
                            "doc_type": document.doc_type,
                            "symbol": document.symbol or "",
                            "sector": document.sector or "",
                            "market": document.market or "",
                            "timestamp": document.timestamp.isoformat(),
                            "path": (document.metadata or {}).get("path"),
                            "user_id": (document.metadata or {}).get("user_id"),
                            "url": (document.metadata or {}).get("url"),
                        }
                    ],
                    ids=[document.doc_id],
                )

            # Store document
            self.documents[document.doc_id] = document

            # Update indexes
            self._update_indexes(document)

            logger.debug(f"Added document to knowledge base: {document.doc_id}")
            return True

        except Exception as e:
            logger.error(f"Failed to add document: {e}")
            return False

    def query_documents(self, query: RAGQuery) -> list[FinancialDocument]:
        """
        Query documents from knowledge base

        Args:
            query: RAG query object

        Returns:
            List[FinancialDocument]: Retrieved documents
        """
        try:
            # Generate query embedding
            query_embedding = self.embedding_system.embed_text(
                query.query_text, query.query_type
            )

            retrieved_docs = []

            if self.collection is not None:
                # Query vector database
                where: dict[str, Any] = {}
                if query.symbols:
                    # Chroma where $in is supported for exact values
                    where["symbol"] = {"$in": query.symbols}
                if query.user_id:
                    where["user_id"] = query.user_id

                try:
                    results = self.collection.query(
                        query_embeddings=[query_embedding.tolist()],
                        n_results=min(query.top_k * 2, 20),
                        include=["documents", "metadatas", "distances"],
                        where=where if where else None,
                    )
                except TypeError:
                    # For older chroma versions without where arg
                    results = self.collection.query(
                        query_embeddings=[query_embedding.tolist()],
                        n_results=min(query.top_k * 2, 20),
                        include=["documents", "metadatas", "distances"],
                    )

                # Convert results to documents
                for i in range(len(results["ids"][0])):
                    doc_id = results["ids"][0][i]
                    if doc_id in self.documents:
                        doc = self.documents[doc_id]
                        doc.relevance_score = (
                            1.0 - results["distances"][0][i]
                        )  # Convert distance to similarity
                        retrieved_docs.append(doc)
            else:
                # Fallback: search in-memory documents
                retrieved_docs = self._fallback_search(query, query_embedding)

            # Apply filters
            filtered_docs = self._apply_filters(retrieved_docs, query)

            # Sort by relevance
            filtered_docs.sort(key=lambda x: x.relevance_score, reverse=True)

            # Backoff: if threshold filtered out everything but we had candidates, relax threshold
            if not filtered_docs and retrieved_docs:
                try:
                    retrieved_docs.sort(key=lambda x: x.relevance_score, reverse=True)
                except Exception:
                    pass
                return retrieved_docs[: query.top_k]

            return filtered_docs[: query.top_k]

        except Exception as e:
            logger.error(f"Document query failed: {e}")
            return []

    def _apply_filters(
        self, documents: list[FinancialDocument], query: RAGQuery
    ) -> list[FinancialDocument]:
        """Apply query filters to documents"""
        filtered = documents

        # Relevance threshold
        filtered = [
            doc for doc in filtered if doc.relevance_score >= query.relevance_threshold
        ]

        # Symbol filter
        if query.symbols:
            filtered = [
                doc
                for doc in filtered
                if doc.symbol is None or doc.symbol in query.symbols
            ]

        # User filter
        if query.user_id:
            filtered = [
                doc
                for doc in filtered
                if (doc.metadata or {}).get("user_id") in {query.user_id}
            ]

        # Document type filter
        if query.doc_types:
            filtered = [doc for doc in filtered if doc.doc_type in query.doc_types]

        # Date range filter
        if query.date_range:
            start_date, end_date = query.date_range
            filtered = [
                doc for doc in filtered if start_date <= doc.timestamp <= end_date
            ]

        return filtered

    def _fallback_search(
        self, query: RAGQuery, query_embedding: np.ndarray
    ) -> list[FinancialDocument]:
        """Fallback search method when vector DB is not available"""
        results = []

        for doc in self.documents.values():
            if doc.embedding is not None:
                similarity = self.embedding_system.compute_similarity(
                    query_embedding, doc.embedding
                )
                doc.relevance_score = similarity
                results.append(doc)

        return results

    def _update_indexes(self, document: FinancialDocument):
        """Update in-memory indexes"""
        doc_id = document.doc_id

        # Symbol index
        if document.symbol:
            self.symbol_index[document.symbol].append(doc_id)

        # Date index (by day)
        date_key = document.timestamp.date().isoformat()
        self.date_index[date_key].append(doc_id)

        # Type index
        self.type_index[document.doc_type].append(doc_id)

    def ingest_news_data(self, symbol: str, days_back: int = 30) -> int:
        """
        Ingest news data for a symbol

        Args:
            symbol: Stock symbol
            days_back: Number of days to look back

        Returns:
            int: Number of documents added
        """
        try:
            news_data = get_finnhub_news(symbol, days_back=days_back)
            added_count = 0

            for news_item in news_data:
                doc_id = f"news_{symbol}_{news_item.get('id', hash(news_item.get('summary', '')))}"

                # Skip if already exists
                if doc_id in self.documents:
                    continue

                document = FinancialDocument(
                    doc_id=doc_id,
                    title=news_item.get("headline", "No Title"),
                    content=news_item.get("summary", ""),
                    doc_type="news",
                    symbol=symbol,
                    timestamp=datetime.fromtimestamp(news_item.get("datetime", 0)),
                    metadata={
                        "source": news_item.get("source", ""),
                        "url": news_item.get("url", ""),
                        "category": news_item.get("category", ""),
                    },
                )

                if self.add_document(document):
                    added_count += 1

            logger.info(f"Ingested {added_count} news documents for {symbol}")
            return added_count

        except Exception as e:
            logger.error(f"Failed to ingest news data for {symbol}: {e}")
            return 0

    def ingest_market_data(self, symbol: str, start_date: str, end_date: str) -> int:
        """
        Ingest market data for technical analysis

        Args:
            symbol: Stock symbol
            start_date: Start date (YYYY-MM-DD)
            end_date: End date (YYYY-MM-DD)

        Returns:
            int: Number of documents added
        """
        try:
            market_data = get_YFin_data_window(symbol, start_date, end_date)

            if market_data is None or market_data.empty:
                return 0

            added_count = 0

            # Create summary documents for different time periods
            periods = ["1D", "1W", "1M"]  # Daily, Weekly, Monthly summaries

            for period in periods:
                if period == "1D":
                    grouped_data = market_data.tail(1)  # Last day
                elif period == "1W":
                    grouped_data = market_data.tail(5)  # Last week
                else:  # 1M
                    grouped_data = market_data.tail(22)  # Last month

                if grouped_data.empty:
                    continue

                # Create technical summary
                summary = self._create_technical_summary(symbol, grouped_data, period)

                doc_id = f"market_{symbol}_{period}_{end_date}"

                document = FinancialDocument(
                    doc_id=doc_id,
                    title=f"{symbol} {period} Technical Summary",
                    content=summary,
                    doc_type="market_data",
                    symbol=symbol,
                    timestamp=datetime.now(),
                    metadata={
                        "period": period,
                        "data_points": len(grouped_data),
                        "price_range": {
                            "high": float(grouped_data["High"].max()),
                            "low": float(grouped_data["Low"].min()),
                            "close": float(grouped_data["Close"].iloc[-1]),
                        },
                    },
                )

                if self.add_document(document):
                    added_count += 1

            logger.info(f"Ingested {added_count} market data documents for {symbol}")
            return added_count

        except Exception as e:
            logger.error(f"Failed to ingest market data for {symbol}: {e}")
            return 0

    def _create_technical_summary(
        self, symbol: str, data: pd.DataFrame, period: str
    ) -> str:
        """Create technical analysis summary from market data"""
        try:
            latest = data.iloc[-1]
            first = data.iloc[0]

            # Basic metrics
            change = latest["Close"] - first["Close"]
            change_pct = (change / first["Close"]) * 100

            # Volume analysis
            avg_volume = data["Volume"].mean()
            latest_volume = latest["Volume"]
            volume_ratio = latest_volume / avg_volume if avg_volume > 0 else 1

            # Price action
            high = data["High"].max()
            low = data["Low"].min()
            volatility = ((high - low) / latest["Close"]) * 100

            summary = f"""
{symbol} {period} Technical Analysis:

Price Action:
- Current Price: ${latest['Close']:.2f}
- Change: ${change:.2f} ({change_pct:.2f}%)
- Range: ${low:.2f} - ${high:.2f}
- Volatility: {volatility:.2f}%

Volume Analysis:
- Latest Volume: {latest_volume:,.0f}
- Average Volume: {avg_volume:,.0f}
- Volume Ratio: {volume_ratio:.2f}x

Market Sentiment: {'Bullish' if change > 0 else 'Bearish' if change < 0 else 'Neutral'}
Period: {period}
Data Points: {len(data)}
"""

            return summary.strip()

        except Exception as e:
            logger.error(f"Failed to create technical summary: {e}")
            return f"{symbol} technical data for {period}"

    def _load_knowledge_base(self):
        """Load existing knowledge base from disk"""
        try:
            kb_file = self.storage_path / "knowledge_base.pkl"
            if kb_file.exists():
                with open(kb_file, "rb") as f:
                    data = pickle.load(f)
                    self.documents = data.get("documents", {})
                    self.symbol_index = data.get("symbol_index", defaultdict(list))
                    self.date_index = data.get("date_index", defaultdict(list))
                    self.type_index = data.get("type_index", defaultdict(list))

                logger.info(
                    f"Loaded {len(self.documents)} documents from knowledge base"
                )
        except Exception as e:
            logger.warning(f"Failed to load existing knowledge base: {e}")

    def save_knowledge_base(self):
        """Save knowledge base to disk"""
        try:
            kb_file = self.storage_path / "knowledge_base.pkl"
            data = {
                "documents": self.documents,
                "symbol_index": dict(self.symbol_index),
                "date_index": dict(self.date_index),
                "type_index": dict(self.type_index),
            }

            with open(kb_file, "wb") as f:
                pickle.dump(data, f)

            logger.info("Knowledge base saved successfully")
        except Exception as e:
            logger.error(f"Failed to save knowledge base: {e}")

    def get_stats(self) -> dict[str, Any]:
        """Get knowledge base statistics"""
        vectors = None
        try:
            if self.collection is not None and hasattr(self.collection, "count"):
                vectors = int(self.collection.count())
        except Exception:
            vectors = None
        return {
            "total_documents": len(self.documents),
            "total_chunks": len(self.documents),
            "vector_count": vectors,
            "documents_by_type": {
                doc_type: len(doc_ids) for doc_type, doc_ids in self.type_index.items()
            },
            "symbols_covered": len(self.symbol_index),
            "date_range": {
                "earliest": min(self.date_index.keys()) if self.date_index else None,
                "latest": max(self.date_index.keys()) if self.date_index else None,
            },
            "storage_path": str(self.storage_path),
            "vector_db_available": self.collection is not None,
        }


class FinancialRAGSystem:
    """Complete Financial RAG System"""

    def __init__(
        self,
        knowledge_base_path: str = "financial_kb",
        llm_orchestrator=None,
        embedding_config: dict[str, Any] | None = None,
    ):
        self.knowledge_base = FinancialKnowledgeBase(
            knowledge_base_path, embedding_config=embedding_config
        )
        self.llm_orchestrator = llm_orchestrator

        # RAG prompt templates
        self.prompt_templates = {
            "general": """Based on the following financial documents, please answer the question:

Question: {query}

Relevant Information:
{context}

Please provide a comprehensive answer based on the information above. If the information is insufficient, please indicate what additional data would be helpful.""",
            "technical": """As a technical analyst, analyze the following market data and answer the question:

Question: {query}

Market Data and Technical Information:
{context}

Please provide technical analysis insights including price action, trends, support/resistance levels, and trading recommendations if applicable.""",
            "fundamental": """As a fundamental analyst, analyze the following financial information and answer the question:

Question: {query}

Financial Data and Research:
{context}

Please provide fundamental analysis including valuation insights, financial health assessment, and investment recommendations based on the available data.""",
            "news": """Based on the latest news and market developments, please answer the question:

Question: {query}

Recent News and Developments:
{context}

Please provide analysis of how these developments might impact the mentioned stocks or markets, including short-term and long-term implications.""",
            "risk": """As a risk analyst, evaluate the following information and answer the question:

Question: {query}

Risk-Related Information:
{context}

Please provide risk assessment including potential downside scenarios, volatility analysis, and risk mitigation recommendations.""",
        }

        logger.info("Financial RAG System initialized")

    async def query(
        self,
        query_text: str,
        query_type: str = "general",
        symbols: list[str] | None = None,
        agent_role: str = "fundamental_expert",
        **kwargs,
    ) -> RAGResponse:
        """
        Execute RAG query with document retrieval and generation

        Args:
            query_text: Query text
            query_type: Type of query
            symbols: Relevant symbols
            agent_role: Agent role for LLM generation
            **kwargs: Additional query parameters

        Returns:
            RAGResponse: Complete RAG response
        """
        try:
            # Create RAG query
            rag_query = RAGQuery(
                query_text=query_text, query_type=query_type, symbols=symbols, **kwargs
            )

            # Retrieve relevant documents
            retrieved_docs = self.knowledge_base.query_documents(rag_query)

            # If no docs are found, still try to answer via LLM directly (non-RAG fallback)
            if not retrieved_docs:
                logger.warning(
                    f"No relevant documents found for query: {query_text[:50]}... (invoking LLM direct fallback)"
                )
                direct_prompt = (
                    "You are a helpful financial analysis assistant.\n"
                    "The knowledge base has no indexed documents for this question right now.\n"
                    "Answer using your general knowledge and reasoning.\n\n"
                    f"Question: {query_text}\n\n"
                    "Provide a clear, concise answer with any useful context."
                )
                generated_response = None
                if self.llm_orchestrator:
                    try:
                        result = await self.llm_orchestrator.execute_task(
                            agent_role=agent_role,
                            task_prompt=direct_prompt,
                            task_type=query_type,
                            context={
                                "rag_enhanced": False,
                                "symbols": symbols,
                                "doc_count": 0,
                                "query_type": query_type,
                                **(kwargs.get("context", {}) or {}),
                            },
                        )
                        if result.success:
                            generated_response = result.result
                    except Exception as e:
                        logger.error(f"Direct LLM fallback failed: {e}")

                if not generated_response:
                    generated_response = (
                        "I can’t find relevant items in the knowledge base yet, "
                        "but here is a general analysis based on the question:\n\n"
                        f"{query_text}"
                    )

                return RAGResponse(
                    query=rag_query,
                    retrieved_documents=[],
                    generated_response=generated_response,
                    confidence_score=0.5,
                    sources=[],
                    metadata={
                        "no_documents_found": True,
                        "llm_direct_fallback": True,
                    },
                )

            # Prepare context from retrieved documents
            context = self._prepare_context(retrieved_docs)

            # Generate response using LLM
            if self.llm_orchestrator:
                generated_response = await self._generate_response(
                    rag_query, context, agent_role, doc_count=len(retrieved_docs)
                )
            else:
                generated_response = self._fallback_response(rag_query, context)

            # Calculate confidence score
            confidence_score = self._calculate_confidence(
                retrieved_docs, generated_response
            )

            # Extract sources
            sources = [f"{doc.title} ({doc.doc_type})" for doc in retrieved_docs]

            return RAGResponse(
                query=rag_query,
                retrieved_documents=retrieved_docs,
                generated_response=generated_response,
                confidence_score=confidence_score,
                sources=sources,
                metadata={
                    "num_documents_retrieved": len(retrieved_docs),
                    "doc_count": len(retrieved_docs),
                    "avg_relevance_score": np.mean(
                        [doc.relevance_score for doc in retrieved_docs]
                    ),
                    "context_length": len(context),
                },
            )

        except Exception as e:
            logger.error(f"RAG query failed: {e}")
            return RAGResponse(
                query=RAGQuery(query_text=query_text, query_type=query_type),
                retrieved_documents=[],
                generated_response=f"I apologize, but I encountered an error while processing your query: {str(e)}",
                confidence_score=0.0,
                sources=[],
                metadata={"error": str(e)},
            )

    # ---------------- Library ingestion (files -> KB) ----------------
    def _read_file_text(self, path: Path) -> tuple[str | None, str | None]:
        """Read supported file types into plain text.

        Returns (text, warn). Warn is a short message when downgraded parsing happens.
        """
        try:
            ext = path.suffix.lower()
            if ext in {".txt", ".md"}:
                return path.read_text(encoding="utf-8", errors="ignore"), None
            if ext in {".csv"}:
                # Lightweight CSV preview: join first ~200 lines
                lines = []
                with open(path, "r", encoding="utf-8", errors="ignore") as f:
                    for i, line in enumerate(f):
                        if i > 2000:
                            lines.append("...")
                            break
                        lines.append(line.rstrip("\n"))
                return "\n".join(lines), "csv_preview"
            if ext in {".html", ".htm"}:
                import re

                raw = path.read_text(encoding="utf-8", errors="ignore")
                # naive tag strip
                txt = re.sub(r"<script[\s\S]*?</script>", " ", raw, flags=re.I)
                txt = re.sub(r"<style[\s\S]*?</style>", " ", txt, flags=re.I)
                txt = re.sub(r"<[^>]+>", " ", txt)
                return txt, "html_stripped"
            if ext in {".docx"}:
                # Prefer python-docx; fallback to pypandoc
                try:
                    from docx import Document  # type: ignore

                    doc = Document(str(path))
                    paras = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
                    return "\n".join(paras), "docx_parsed"
                except Exception:
                    try:
                        import pypandoc  # type: ignore

                        out = pypandoc.convert_file(str(path), to="plain")
                        return out, "docx_pandoc"
                    except Exception:
                        return None, "unsupported_or_parse_failed"
            if ext in {".xlsx"}:
                # Use pandas if available
                try:
                    import pandas as _pd

                    x = _pd.ExcelFile(str(path))
                    parts: list[str] = []
                    for name in x.sheet_names[:5]:  # cap sheets for safety
                        df = x.parse(name)
                        # limit rows/cols for readability
                        df = df.iloc[:200, :20]
                        parts.append(f"### Sheet: {name}\n{df.to_csv(index=False)}")
                    return "\n\n".join(parts), "xlsx_parsed"
                except Exception:
                    return None, "unsupported_or_parse_failed"
            if ext in {".pdf"}:
                # Prefer PyMuPDF; fallback to pandoc
                try:
                    import fitz  # type: ignore

                    text_parts: list[str] = []
                    with fitz.open(str(path)) as doc:
                        for i, page in enumerate(doc):
                            if i >= 50:  # cap pages
                                text_parts.append("...")
                                break
                            text_parts.append(page.get_text())
                    return "\n".join(text_parts), "pdf_extracted"
                except Exception:
                    try:
                        import pypandoc  # type: ignore

                        out = pypandoc.convert_file(str(path), to="plain")
                        return out, "pdf_pandoc"
                    except Exception:
                        return None, "unsupported_or_parse_failed"
        except Exception:
            pass
        return None, "unsupported_or_parse_failed"

    def _chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> list[str]:
        if chunk_size <= 0:
            chunk_size = 1000
        if overlap < 0:
            overlap = 0
        res: list[str] = []
        n = len(text)
        i = 0
        while i < n:
            j = min(i + chunk_size, n)
            seg = text[i:j].strip()
            if seg:
                res.append(seg)
            if j >= n:
                break
            i = j - overlap if overlap > 0 else j
            if i < 0:
                i = 0
        return res

    def ingest_library(
        self,
        root_dir: str,
        user_id: str | None = None,
        symbol: str | None = None,
        chunk_size: int = 1000,
        overlap: int = 200,
        files: list[str] | None = None,
        dry_run: bool = False,
    ) -> dict[str, Any]:
        """Ingest files under root_dir into the knowledge base.

        Minimal implementation: supports txt/md/csv/html; others are skipped with warnings.
        Creates chunked FinancialDocuments and persists KB.
        """
        root = Path(root_dir)
        if not root.exists() or not root.is_dir():
            raise ValueError(f"Library root not found: {root}")

        added = 0
        skipped = 0
        warnings: list[str] = []

        exts = {".txt", ".md", ".csv", ".html", ".htm", ".pdf", ".docx", ".xlsx"}

        candidates: list[Path]
        if files:
            candidates = []
            for f in files:
                fp = Path(f)
                if not fp.is_absolute():
                    fp = root / fp
                if fp.exists() and fp.is_file():
                    candidates.append(fp)
        else:
            candidates = [p for p in sorted(root.rglob("*")) if p.is_file()]

        for p in candidates:
            if not p.is_file():
                continue
            if p.suffix.lower() not in exts:
                skipped += 1
                continue
            text, warn = self._read_file_text(p)
            if warn == "unsupported_or_parse_failed" or not text:
                warnings.append(f"Parse skipped: {p}")
                skipped += 1
                continue
            if warn and warn != "unsupported_or_parse_failed":
                warnings.append(f"{warn}: {p}")

            chunks = self._chunk_text(text, chunk_size=chunk_size, overlap=overlap)
            if not chunks:
                skipped += 1
                continue

            # document metadata
            title = p.stem
            doc_type = "library"
            try:
                ts = datetime.fromtimestamp(p.stat().st_mtime)
            except Exception:
                ts = datetime.now()
            # stable file hash
            try:
                h = hashlib.md5(p.read_bytes()).hexdigest()
            except Exception:
                h = hashlib.md5(str(p).encode("utf-8")).hexdigest()

            for idx, seg in enumerate(chunks):
                doc_id = f"lib_{h}_{idx}"
                if doc_id in self.knowledge_base.documents:
                    skipped += 1
                    continue
                if dry_run:
                    added += 1  # count as would-add
                    continue

                # Build file URL when static serving is enabled and the path is under root_dir
                file_url = None
                try:
                    serve_static = str(os.getenv("SERVE_LIBRARY_STATIC", "true")).lower() in {
                        "1",
                        "true",
                        "yes",
                        "on",
                    }
                    if serve_static:
                        rel = p.resolve().relative_to(root.resolve())
                        prefix = os.getenv("LIBRARY_URL_PREFIX", "/library").rstrip("/")
                        rel_url = "/".join(str(rel).split(os.sep))
                        file_url = f"{prefix}/{rel_url}"
                except Exception:
                    file_url = None

                doc = FinancialDocument(
                    doc_id=doc_id,
                    title=title,
                    content=seg,
                    doc_type=doc_type,
                    symbol=symbol,
                    timestamp=ts,
                    metadata={
                        "path": str(p),
                        "ext": p.suffix.lower(),
                        "user_id": user_id,
                        "url": file_url,
                    },
                )
                if self.knowledge_base.add_document(doc):
                    added += 1
                else:
                    skipped += 1

        # Persist KB
        if not dry_run:
            try:
                self.knowledge_base.save_knowledge_base()
            except Exception as e:
                warnings.append(f"save_kb_failed: {e}")

        return {"added": added, "skipped": skipped, "warnings": warnings}

    def _prepare_context(self, documents: list[FinancialDocument]) -> str:
        """Prepare context string from retrieved documents"""
        context_parts = []

        for i, doc in enumerate(documents, 1):
            # Format document information
            doc_context = f"""
Document {i}: {doc.title}
Type: {doc.doc_type.replace('_', ' ').title()}
{f"Symbol: {doc.symbol}" if doc.symbol else ""}
Date: {doc.timestamp.strftime('%Y-%m-%d')}
Relevance: {doc.relevance_score:.2f}

Content:
{doc.content}
"""
            context_parts.append(doc_context.strip())

        return "\n\n" + "\n\n".join(context_parts)

    async def _generate_response(
        self, query: RAGQuery, context: str, agent_role: str, doc_count: int | None = None
    ) -> str:
        """Generate response using LLM orchestrator"""
        try:
            # Select appropriate prompt template
            template = self.prompt_templates.get(
                query.query_type, self.prompt_templates["general"]
            )

            # Format prompt
            prompt = template.format(query=query.query_text, context=context)

            # Execute with LLM orchestrator
            result = await self.llm_orchestrator.execute_task(
                agent_role=agent_role,
                task_prompt=prompt,
                task_type=query.query_type,
                context={
                    "rag_enhanced": True,
                    "symbols": query.symbols,
                    "doc_count": int(doc_count or 0),
                    "query_type": query.query_type,
                },
            )

            if result.success:
                return result.result
            else:
                return f"I encountered an issue generating a response: {result.error_message}"

        except Exception as e:
            logger.error(f"LLM response generation failed: {e}")
            return f"I apologize, but I couldn't generate a proper response due to a technical issue: {str(e)}"

    def _fallback_response(self, query: RAGQuery, context: str) -> str:
        """Fallback response when LLM orchestrator is not available"""
        return f"""Based on the available financial documents, here's the relevant information for your query: "{query.query_text}"

{context}

Note: This is a document-based response. For more detailed analysis, please ensure the AI analysis system is properly configured."""

    def _calculate_confidence(
        self, documents: list[FinancialDocument], response: str
    ) -> float:
        """Calculate confidence score for the response"""
        if not documents:
            return 0.0

        # Base confidence from document relevance
        avg_relevance = np.mean([doc.relevance_score for doc in documents])

        # Adjust based on response quality indicators
        response_lower = response.lower()

        # Positive indicators
        confidence_boost = 0.0
        if any(
            phrase in response_lower
            for phrase in ["based on", "according to", "data shows"]
        ):
            confidence_boost += 0.1

        if len(response) > 200:  # Detailed response
            confidence_boost += 0.05

        # Negative indicators
        confidence_penalty = 0.0
        if any(
            phrase in response_lower
            for phrase in ["insufficient", "unclear", "unable to"]
        ):
            confidence_penalty += 0.2

        # Calculate final confidence
        confidence = avg_relevance + confidence_boost - confidence_penalty
        return max(0.0, min(1.0, confidence))  # Clamp to [0, 1]

    def ingest_symbol_data(
        self,
        symbol: str,
        include_news: bool = True,
        include_market_data: bool = True,
        days_back: int = 30,
    ) -> dict[str, int]:
        """
        Ingest comprehensive data for a symbol

        Args:
            symbol: Stock symbol
            include_news: Whether to include news data
            include_market_data: Whether to include market data
            days_back: Number of days to look back

        Returns:
            Dict[str, int]: Ingestion statistics
        """
        stats = {"news": 0, "market_data": 0}

        if include_news:
            stats["news"] = self.knowledge_base.ingest_news_data(symbol, days_back)

        if include_market_data:
            end_date = datetime.now().strftime("%Y-%m-%d")
            start_date = (datetime.now() - timedelta(days=days_back)).strftime(
                "%Y-%m-%d"
            )
            stats["market_data"] = self.knowledge_base.ingest_market_data(
                symbol, start_date, end_date
            )

        # Save knowledge base after ingestion
        self.knowledge_base.save_knowledge_base()

        logger.info(f"Ingested data for {symbol}: {stats}")
        return stats

    def _detect_parsers(self) -> dict[str, bool]:
        """Detect optional parsing dependencies for PDF/DOCX/XLSX and pandoc."""
        out = {
            "pypandoc": False,
            "pandoc_binary": False,
            "python_docx": False,
            "pymupdf": False,
            "pandas": False,
            "openpyxl": False,
        }
        try:
            import pypandoc  # type: ignore  # noqa: F401

            out["pypandoc"] = True
            try:
                import shutil, subprocess  # noqa: E401

                if shutil.which("pandoc"):
                    subprocess.run(["pandoc", "--version"], capture_output=True, timeout=3)
                    out["pandoc_binary"] = True
            except Exception:
                pass
        except Exception:
            pass
        try:
            import docx  # type: ignore  # noqa: F401

            out["python_docx"] = True
        except Exception:
            pass
        try:
            import fitz  # type: ignore  # noqa: F401

            out["pymupdf"] = True
        except Exception:
            pass
        try:
            import pandas  # noqa: F401

            out["pandas"] = True
            try:
                import openpyxl  # noqa: F401

                out["openpyxl"] = True
            except Exception:
                pass
        except Exception:
            pass
        return out

    def get_system_stats(self) -> dict[str, Any]:
        """Get comprehensive system statistics including parser detection."""
        kb_stats = self.knowledge_base.get_stats()
        serve_static = str(os.getenv("SERVE_LIBRARY_STATIC", "true")).lower() in {
            "1",
            "true",
            "yes",
            "on",
        }
        emb_avail = False
        try:
            emb = self.knowledge_base.embedding_system
            emb_avail = bool(getattr(emb, "siliconflow_api_key", None)) or bool(
                getattr(emb, "model", None)
            )
        except Exception:
            emb_avail = False

        return {
            "knowledge_base": kb_stats,
            "rag_system": {
                "prompt_templates": len(self.prompt_templates),
                "llm_orchestrator_available": self.llm_orchestrator is not None,
                "embedding_system_available": emb_avail,
                "vector_db_available": self.knowledge_base.collection is not None,
                "parsers": self._detect_parsers(),
                "library_root": os.getenv("LIBRARY_ROOT", "./data/library"),
                "library_url_prefix": os.getenv("LIBRARY_URL_PREFIX", "/library"),
                "serve_library_static": serve_static,
            },
        }

    # =============== Library ingestion (MVP, obsolete – kept for backward-compat) ===============
    def parse_file_to_text_obsolete(self, path: str | Path) -> str | None:
        """Parse file to plain text (MVP): supports .txt/.md/.csv/.html minimal.

        For other formats, return None to skip (keeps compatibility minimal).
        """
        try:
            p = Path(path)
            if not p.exists() or not p.is_file():
                return None
            ext = p.suffix.lower()
            if ext in {".txt", ".md"}:
                return p.read_text(encoding="utf-8", errors="ignore")
            if ext == ".csv":
                import pandas as _pd

                try:
                    df = _pd.read_csv(p)
                    return df.to_csv(index=False)
                except Exception:
                    return p.read_text(encoding="utf-8", errors="ignore")
            if ext in {".html", ".htm"}:
                # strip tags minimal
                import re as _re

                html = p.read_text(encoding="utf-8", errors="ignore")
                text = _re.sub(r"<[^>]+>", " ", html)
                return text
            # unsupported
            return None
        except Exception as e:
            logger.warning(f"parse_file_to_text failed for {path}: {e}")
            return None

    def chunk_text_obsolete(
        self, text: str, max_len: int = 1000, overlap: int = 100
    ) -> list[str]:
        """Simple text chunking by characters with overlap (MVP)."""
        chunks: list[str] = []
        n = len(text or "")
        if n == 0:
            return chunks
        start = 0
        while start < n:
            end = min(n, start + max_len)
            chunks.append(text[start:end])
            if end == n:
                break
            start = max(end - overlap, start + 1)
        return chunks

    def ingest_library_obsolete(
        self,
        root_dir: str | Path,
        default_doc_type: str = "library",
        user_id: str | None = None,
        symbol: str | None = None,
    ) -> dict[str, Any]:
        """Ingest a directory of plain files into the knowledge base (MVP).

        Only processes lightweight formats (.txt/.md/.csv/.html). Others are skipped safely.
        """
        root = Path(root_dir)
        if not root.exists() or not root.is_dir():
            return {"added": 0, "skipped": 0, "warnings": [f"Invalid root_dir: {root}"]}

        added = 0
        skipped = 0
        warnings: list[str] = []

        for p in root.rglob("*"):
            if not p.is_file():
                continue
            text = self.parse_file_to_text(p)
            if not text:
                skipped += 1
                continue
            chunks = self.chunk_text(text)
            base_id = hashlib.md5(str(p.resolve()).encode()).hexdigest()
            for i, chunk in enumerate(chunks):
                doc = FinancialDocument(
                    doc_id=f"lib_{base_id}_{i}",
                    title=p.stem,
                    content=chunk,
                    doc_type=default_doc_type,
                    symbol=symbol,
                    metadata={
                        "source_path": str(p),
                        "user_id": user_id or "",
                    },
                )
                if self.knowledge_base.add_document(doc):
                    added += 1
                else:
                    warnings.append(f"Add failed: {p}#{i}")
        # Persist
        try:
            self.knowledge_base.save_knowledge_base()
        except Exception as e:
            warnings.append(f"save_knowledge_base failed: {e}")

        return {"added": added, "skipped": skipped, "warnings": warnings}
